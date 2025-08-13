from azure.core.credentials import AzureKeyCredential
from azure.ai.formrecognizer import DocumentAnalysisClient, AnalyzeResult
from app.core.config import AZURE_ENDPOINT, AZURE_KEY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_API_VERSION,AZURE_DEPLOYMENT_NAME,AZURE_OPENAI_KEY, AZURE_STORAGE_CONNECTION_STRING, CONTAINER_NAME,AZURE_STORAGE_ACCOUNT_KEY
from app.model.models import ApiResponse
from openai import AzureOpenAI
from app.api.endpoints.chat.function_call import Available_functions,validate_function_call,get_function_by_name,FunctionNames
import uuid,os,asyncio,urllib,json
from azure.storage.blob import BlobServiceClient,generate_blob_sas, BlobSasPermissions
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
from datetime import datetime, timedelta,timezone
from langdetect import detect
from tempfile import NamedTemporaryFile
import pandas as pd
from fastapi.responses import JSONResponse
from openpyxl.styles import Font, PatternFill, Alignment




try:
    document_client = DocumentAnalysisClient(
        endpoint=AZURE_ENDPOINT,
        credential=AzureKeyCredential(AZURE_KEY)
    )
    azure_client = AzureOpenAI(
        azure_endpoint = AZURE_OPENAI_ENDPOINT,
        api_key = AZURE_OPENAI_KEY,
        api_version = AZURE_OPENAI_API_VERSION
    )
except Exception as e:
    raise



system_prompt = """
You are an intelligent document analysis assistant. Your primary role is to help users analyze PDF documents by extracting text and tables, then saving the results.

CORE CAPABILITIES:
1. Extract handwritten and typed text from PDF documents
2. Extract and analyze table data from PDFs  

FUNCTION CALLING RULES:
When a user uploads a PDF and their message contains ANY of these keywords or phrases, you MUST call the appropriate function:

FOR TEXT EXTRACTION - Call "extract_handwritten_from_document":
- Keywords: content, text, handwritten, handwriting, typed, document
- Phrases: "extract text", "get content", "read document", "pull text", "what does it say", "analyze document"
- Even variations like: "extract the typed part", "pull handwritten notes", "get all text"

FOR TABLE EXTRACTION - Call "extract_table_from_document":
- Keywords: table, tables, data, tabular, rows, columns, spreadsheet, grid, structured data
- Phrases: "extract tables", "get table data", "analyze tables"

FOLLOW-UP BEHAVIOR LOGIC:
1. If the user **asks to extract text or tables but doesn’t upload a file in that message**:
   → Check the **last 5 messages** in the conversation for a previously uploaded PDF
   → If a PDF is found, call the appropriate function using the last uploaded file
   → If no PDF is found, respond: **"Please upload a PDF so I can begin extracting the content."**

2. If the user **uploads a file and gives a vague or unclear message** (e.g., "check this", "see attached"):
   → Ask: **"Would you like me to extract the text content or the table data from this document?"**
   → Do not call a function unless the user's intent is clarified.

3. If the user **greets you or sends a general message without asking for extraction**:
   → Respond warmly and say: **"Hi! Please upload a PDF so I can help you extract text or table data."**

4. If the user **uploads a file without any message**:
   → Ask: **"What would you like me to extract from this document – text content, tables, or both?"**

5. If the user **initially asks to extract something**, and **only uploads a file in a follow-up message**, then:
   → Track conversation history (up to last 5 messages)
   → When file is uploaded, immediately trigger the correct function based on the prior request
   → Do NOT ask again for intent if it was already stated earlier.

LANGUAGE DETECTION:
- If the user's message is in **Hungarian**, respond in **Hungarian only**
- If the user switches to English or another language, resume responding in that language
- Maintain behavior and logic regardless of language – only the response language changes

BEHAVIOR GUIDELINES:
- Always call functions when keyword intent is clear – don't wait for confirmation
- Default to text extraction if you’re unsure whether user wants text or table
- Always check for recently uploaded files (up to last 5 messages)
- Be proactive, friendly, and focused only on document-related tasks
- Never answer general knowledge or unrelated questions
- Don't explain technical steps unless explicitly asked

EXAMPLES:
- User: "extract this" + uploads PDF → Call extract_handwritten_from_document
- User: "what's in this document?" + uploads PDF → Call extract_handwritten_from_document  
- User: "get the tables from this" + uploads PDF → Call extract_table_from_document
- User: "extract text" (no file) → Check last 5 messages → If file found, extract text; else ask for upload
- User: uploads file + says "check this" → Ask: "Would you like me to extract the text content or the table data from this document?"
- User: "extract tables" → Assistant says "Please upload a PDF..." → User uploads → Assistant calls extract_table_from_document immediately
"""

system_prompt1 = """
You are an intelligent document analysis assistant.

Your job is to:
1. Analyze delivery notes written in Hungarian (even if some parts are translated).
2. Extract all relevant information regardless of layout or formatting.
3. Organize the extracted data in a top-to-bottom logical order.
4. Format the output using valid and clean HTML tags suitable for UI display.
5. Do not add or generate any new section headings or English translations — always retain the original Hungarian content.
6. Do not write any extra content apart from the structured headings and extracted values.

---------------------------------------------
🧠 STRUCTURE & SEMANTICS (LOGIC ONLY):
---------------------------------------------

- Maintain the original Hungarian structure and section names from the input.
- Do NOT create or insert your own headings (e.g., "Shipping", "Company Details", etc.).
- Use the order and hierarchy implied by the document itself.
- Always output in top-to-bottom order regardless of original formatting (e.g., left-to-right columns).

---------------------------------------------
🎨 HTML OUTPUT FORMAT:
---------------------------------------------

- Use <h2> for major sections (e.g., "Szállítólevél", "Megrendelő", etc.)
- Use <h3> if needed for subsections (e.g., "Áru címzettje")
- Use <p><b>Mező:</b> Érték</p> for individual fields
- Use <ul><li>Mező: Érték</li></ul> for grouped information (e.g., address blocks)
- Use <table> with <thead> and <tbody> for itemized goods or acceptance time tables
- Use <blockquote> for declarations or remarks (e.g., "Nem megfelelő – Minőségi kifogás miatt visszautasítva!")

------------------------------------------------
🛑 STRICT EXCLUSIONS:
------------------------------------------------

- ❌ Do NOT generate your own English or Hungarian headings — use only what is present in the original content
- ❌ Do NOT include summaries, interpretations, or explanations
- ❌ Do NOT repeat the raw input text
- ❌ Do NOT add placeholder fields for missing values
- ❌ Do NOT prepend or append anything like "Here is the result"

------------------------------------------------
📌 LANGUAGE HANDLING:
------------------------------------------------

- Assume all delivery note content is Hungarian
- Do NOT translate, rewrite, or summarize content into English or any other language
- Preserve all Hungarian text, field names, and values as they appear

------------------------------------------------
🎯 GOAL:
------------------------------------------------

Extract and reformat delivery note content into clean, readable, and structured HTML output using only the original Hungarian section names and values.
Ensure the output is suitable for embedding in a UI and includes no extra content or fabricated structure.
"""
CONVERSATION_HISTORY = {}  # Store conversation history with file content by session_id

async def fn_chatbot(message=None, file=None, session_id=None) -> ApiResponse:
    """ Function to handle the chat request with conversation history
    Args:
        message (str): The message to be processed.
        file (UploadFile): The file to be analyzed.
        session_id (str): Session identifier for maintaining conversation history.
    Returns:
        ApiResponse: The response containing the analysis result or OpenAI response.
    """
    try:
        global blob_name, CONVERSATION_HISTORY
        
        if not session_id:
            session_id = str(uuid.uuid4())
            
        # Initialize conversation history for new sessions
        if session_id not in CONVERSATION_HISTORY:
            CONVERSATION_HISTORY[session_id] = {
                'file': None,
                'file_content': None,
                'History': []
            }
            
        if file:
            if not file.filename.endswith('.pdf') or not file.content_type.startswith('application/pdf'):
                raise ValueError("File must be a PDF document")

            # Store file content in conversation history instead of temp store
            file_content = await file.read()
            CONVERSATION_HISTORY[session_id]['file_content'] = file_content
            CONVERSATION_HISTORY[session_id]['file'] = file.filename
            if "handwritten" in message.lower() or "hand written" in message.lower() or "kézzel írott kivonat" in message or "kézzel írott" in message:
                blob_name = f"{uuid.uuid4()}_{file.filename.replace('.pdf','')}.docx"

            for entry in reversed(CONVERSATION_HISTORY[session_id]['History']):
                    if entry['role'] == 'user':
                        last_user_message = entry['content']
                        break

            try:
                language = detect(last_user_message)
            except:
                language = "en"  # fallback in case of detection failure

            if language == "hu":
                message = "a fájl feltöltve a felhasználó által" if message == "Empty Message" else f"{message} és a fájl is feltöltve van a felhasználó által"
            else:
                message = "file is uploaded by user" if message == "Empty Message" else f"{message} and also file is uploaded by user"
            # message = "file is uploaded by user" if message == "Empty Message" else f"{message} and also file is uploaded by user"
            

        
            
        # Get file content from conversation history
        file_content = CONVERSATION_HISTORY[session_id].get('file_content')
        response = await fn_chat_with_bot(message, file_content, session_id)
        # if type(response) == dict:
        #     response['session_id'] = session_id
        
        # return ApiResponse(
        #     statusCode=200,
        #     message="chatbot response",
        #     data=response
        # )
        # if "table" in message.lower():
        if isinstance(response, dict):
            response['session_id'] = session_id
        elif isinstance(response, ApiResponse):
            if not isinstance(response.data, dict):
                    response.data = {}
            response.data['session_id'] = session_id
        return response
        # else:
        #     if type(response) == dict:
        #         response['session_id'] = session_id
        
        #     return ApiResponse(
        #         statusCode=200,
        #         message="chatbot response",
        #         data=response
        #     )
    except Exception as e:
        return ApiResponse(
            statusCode=500,
            message=str(e)
        )

async def fn_chat_with_bot(message, file_content, session_id):
    """
    Function to interact with the Azure OpenAI with conversation history
    Args:
        message (str): The message to be processed.
        file_content (bytes): The file content to be analyzed.
        session_id (str): Session identifier for conversation history.
    Returns:
        dict: The response containing the analysis result or OpenAI response.
    """
    try:
        global CONVERSATION_HISTORY
        
        # Only add user message if not already added in initial greeting handler
        if not (CONVERSATION_HISTORY[session_id]['History'] and 
                CONVERSATION_HISTORY[session_id]['History'][-1].get('role') == 'user' and 
                CONVERSATION_HISTORY[session_id]['History'][-1].get('content') == message):
            CONVERSATION_HISTORY[session_id]['History'].append({"role": "user", "content": message})
        
        # Prepare messages with history
        messages = [{"role": "system", "content": system_prompt}]
        
        # Add conversation history (limit to last 20 exchanges to manage token usage)
        history_limit = 5  # 10 user + 10 assistant messages
        recent_history = CONVERSATION_HISTORY[session_id]['History'][-history_limit:] if len(CONVERSATION_HISTORY[session_id]['History']) > history_limit else CONVERSATION_HISTORY[session_id]['History']
        messages.extend(recent_history)
        
        response = azure_client.chat.completions.create(
            model=AZURE_DEPLOYMENT_NAME,
            messages=messages,
            tools=Available_functions,
            tool_choice="auto"
        )
        
        response_message = response.choices[0].message
        
        # Handle function calls
        if hasattr(response_message, 'tool_calls') and response_message.tool_calls:
            function_name = response_message.tool_calls[0].function.name
            function_args = json.loads(response_message.tool_calls[0].function.arguments)
            
            if not validate_function_call(function_name, function_args, file_content):
                error_response = "please Upload a Document."
                for entry in reversed(CONVERSATION_HISTORY[session_id]['History']):
                    if entry['role'] == 'user':
                        last_user_message = entry['content']
                        break

                try:
                    language = detect(last_user_message)
                except:
                    language = "en"  # fallback in case of detection failure

                if language == "hu":
                    error_response = "Kérjük, töltsön fel egy dokumentumot."
                else:
                    error_response = "Please upload a document."

                CONVERSATION_HISTORY[session_id]['History'].append({"role": "assistant", "content": error_response})
                return ApiResponse(statusCode=400, message=error_response, data={})
            
            function_response = await call_function(function_name, function_args, file_content, session_id)
            
            # Add success/failure status to history instead of full response
            status_message = f"Function '{function_name}' executed successfully" if function_response else f"Function '{function_name}' failed to execute"
            CONVERSATION_HISTORY[session_id]['History'].append({"role": "assistant", "content": status_message})
            
            return ApiResponse(
                statusCode=function_response.statusCode,
                message=function_response.message,
                data=function_response.data
            )
        
        # Regular response without function calls
        assistant_response = response_message.content
        CONVERSATION_HISTORY[session_id]['History'].append({"role": "assistant", "content": assistant_response})
        
        return ApiResponse(
            statusCode=200,
            message="Chatbot response",
            data={"reply": assistant_response}
        )
        
    except Exception as e:
        error_message = f"Error in chatbot functions: {str(e)}"
        CONVERSATION_HISTORY[session_id]['History'].append({"role": "assistant", "content": error_message})
        return ApiResponse(
            statusCode=500,
            message=f"Error in chatbot functions: {str(e)}",
            data={}
        )



async def call_function(function_name, function_args, file,session_id):
    """
    call the function based on the function name and arguments
    args:
    function_name (str): The name of the function to call.
    function_args (dict): The arguments to pass to the function.
    file (uploadFile): The file to be analyzed.
    returns:
    dict: The response from the function call.
    """
    function = get_function_by_name(function_name)
    if not function:
        return ApiResponse(statusCode=500, message=f"Function {function_name} not found", data={})
    
    if function['name'] == FunctionNames.EXTRACT_HANDWRITING:
        result = await extract_handwritten_from_document(session_id,file)
        return ApiResponse(
        statusCode=200,
        message="Handwritten extraction successful",
        data=result
       )
    elif function['name'] == FunctionNames.EXTRACT_TABLE:
        result = await fn_analyze_document(file,session_id)
        CONVERSATION_HISTORY[session_id]['file_content'] = None
        return result
        
    else:
        if detect_language_from_user_history(session_id) == 'en':
           return ApiResponse(statusCode=500, message=f"Unknown Function {function_name}",data={})
        else:
            return ApiResponse(statusCode=500, message=f"Ismeretlen függvény: {function_name}",data={})  
    

async def extract_handwritten_from_document(session_id,file):
    """Extract handwritten text from a document"""
    try:
        if not file:
            raise ValueError("File is required for extraction")
        
        # file_bytes =await file.read()
        poller = document_client.begin_analyze_document(
            model_id = "prebuilt-layout",
            document = file
        )

        result: AnalyzeResult = poller.result()

        if result.pages:
            # print("--- Document Analysis Results (Read Model) ---")
            
            # Print the full document content for context
            # print(f"Full document content:\n{result.content}\n")

            # Iterate through each page
            for page in result.pages:
                # print(f"==================== Page {page.page_number} ====================")
                
                # Iterate through each line on the page
                for line in page.lines:
                    
                    # Determine if the line is handwritten
                    is_handwritten = False
                    if result.styles:
                        for style in result.styles:
                            # Find the style that corresponds to this line's text
                            if style.is_handwritten:
                                for span in line.spans:
                                    if span.offset >= style.spans[0].offset and span.offset + span.length <= style.spans[0].offset + style.spans[0].length:
                                        is_handwritten = True
                                        break
                                if is_handwritten:
                                    break

                    line_type = "Handwritten" if is_handwritten else "Printed"
                    
                    # print(f"Line ({line_type}): '{line.content}'")

                    # Now, let's get the confidence score for each word in the line
                    line_words = []
                    for word in page.words:
                        # Check if the word's bounding box is within the line's bounding box
                        # A more robust check might involve span offsets, but this is often sufficient
                        if (word.polygon and line.polygon and 
                            word.polygon[0].y >= line.polygon[0].y and 
                            word.polygon[2].y <= line.polygon[2].y):
                            
                            # Get the confidence score for each word
                            line_words.append(f"'{word.content}' ({word.confidence:.2f})")
                    
                    # Print the words and their confidence scores for the current line
                    # if line_words:
                        # print(f"  Words & Confidence: {', '.join(line_words)}")
                
                # print("\n")
            Total_confidence = sum(word.confidence for page in result.pages for word in page.words) / max(1, sum(len(page.words) for page in result.pages))
            
            # print("--- End of Analysis ---")
            
        else:
            print("No pages or text found in the document.")

        # blob_url = None
        # blob_url = await save_extracted_handwritten_text(result.content)

        CONVERSATION_HISTORY[session_id]['file_content'] = None
        response = azure_client.chat.completions.create(
            model = AZURE_DEPLOYMENT_NAME,
            messages=[{"role": "system", "content": system_prompt1},
                      {"role":"user", "content": result.content}]
        )
        response_message = response.choices[0].message.content

        blob_url = None
        blob_url = await save_extracted_handwritten_text(response_message)

        return {
                "response": response_message,
                "confidence": Total_confidence,
                "blob_url": blob_url
                }
        
    except Exception as e:
        raise Exception(f"Error extracting handwritten text: {str(e)}")

async def save_extracted_handwritten_text(extracted_text):
    """
    Save the extracted handwritten text to a storage
    """
    try:
        structured_text = await html_to_structured_text(extracted_text)

        # Convert to Word document
        doc = await structured_text_to_docx(structured_text)
        word_stream = BytesIO()
        doc.save(word_stream)
        word_stream.seek(0)
        
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        blob_client = blob_service_client.get_blob_client(container=CONTAINER_NAME, blob=blob_name)


        blob_client.upload_blob(word_stream, overwrite=True)
        print(f"Uploaded '{blob_name}' to container '{CONTAINER_NAME}' successfully.")


        blob_url = blob_client.url
        return blob_url
    except Exception as e:
        raise Exception(f"Error saving extracted handwritten text: {str(e)}")

async def fn_save_text(blob_url):
    """ Function to save the extracted text to a storage permanetly
    Args:
        blob_url (str): The temporary blob URL of the extracted text.
    Returns:
        ApiResponse: The response containing the permanent blob URL with status code and success message.
    """
    try:
        if not blob_url:
            raise ValueError("Blob URL is required to save the text")
        
        Permanent_container_name = "extracted-files/files"
        blob_name = urllib.parse.unquote(blob_url.split('/')[-1])
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        blob_client = blob_service_client.get_blob_client(container=Permanent_container_name, blob=blob_name)
        source_blob_url = blob_url
        blob_client.start_copy_from_url(source_blob_url)

        props = blob_client.get_blob_properties()
        while props.copy.status == "pending":
            await asyncio.sleep(1)
            props = blob_client.get_blob_properties()
        download_sas_url = generate_download_sas_url(blob_name, Permanent_container_name)
        await delete_blob(blob_name)
        return ApiResponse(
            statusCode=200,
            message="Text saved successfully",
            data={"response": "download the file from the permanent url", "permanent_blob_url": download_sas_url}
        )
    except Exception as e:
        return ApiResponse(
            statusCode=500,
            message=str(e)
        )

def generate_download_sas_url(blob_name: str, container_name: str) -> str:
    """Generate a 1-hour download URL (SAS URL) for a blob."""
    blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
    sas_token = generate_blob_sas(
        account_name=blob_service_client.account_name,
        container_name=container_name,
        blob_name=blob_name,
        account_key=blob_service_client.credential.account_key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(hours=1)
    )

    sas_url = f"https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/{urllib.parse.quote(blob_name)}?{sas_token}"
    return sas_url

async def delete_blob(blob_name):
    try:
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        blob_client = blob_service_client.get_blob_client(container=CONTAINER_NAME, blob=blob_name)

        blob_client.delete_blob()
    except Exception as e:
        raise Exception(f"Error saving extracted handwritten text: {str(e)}")





async def html_to_structured_text(text):
    soup = BeautifulSoup(text, 'html.parser')
    output = []

    def clean_text(t):
        return ' '.join(t.strip().split())

    def handle_element(el):
        if el.name == 'h2':
            output.append(f"\n### **{clean_text(el.get_text())}**")

        elif el.name == 'ul':
            for li in el.find_all('li'):
                li_text = clean_text(li.get_text())
                if ':' in li_text:
                    parts = li_text.split(':', 1)
                    output.append(f"- **{parts[0].strip()}:** {parts[1].strip()}")
                else:
                    output.append(f"- {li_text}")
            output.append("")

        elif el.name == 'p':
            bold_tags = el.find_all('b')
            if bold_tags:
                text_parts = []
                for content in el.contents:
                    if content.name == 'b':
                        text_parts.append(f"**{clean_text(content.get_text())}**")
                    elif content.string:
                        text_parts.append(clean_text(content.string))
                output.append(' '.join(text_parts))
            else:
                output.append(clean_text(el.get_text()))

        elif el.name == 'table':
            thead = el.find('thead')
            tbody = el.find('tbody')
            if not thead or not tbody:
                return
            headers = [clean_text(th.get_text()) for th in thead.find_all('th')]
            output.append('\n' + ' | '.join(headers))
            output.append('|'.join(['---'] * len(headers)))
            for tr in tbody.find_all('tr'):
                row = [clean_text(td.get_text()) for td in tr.find_all('td')]
                output.append(' | '.join(row))
            output.append("")

        elif el.name == 'blockquote':
            output.append(f"> {clean_text(el.get_text())}")

    # Safely traverse top-level elements in correct order
    top_level_elements = soup.body.find_all(recursive=False) if soup.body else soup.find_all(recursive=False)
    for element in top_level_elements:
        handle_element(element)

    return '\n'.join(output)



        

async def structured_text_to_docx(structured_text: str) -> Document:
    doc = Document()
    lines = structured_text.splitlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Header (### **Header**)
        if line.startswith("### **") and line.endswith("**"):
            heading = line.replace("### **", "").rstrip("**")
            doc.add_heading(heading.strip(), level=2)

        # Bullet point with bold label: "- **Label:** Value"
        elif line.startswith("- **") and ":**" in line:
            parts = line[4:].split(":**", 1)
            label = parts[0].strip()
            value = parts[1].strip()
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(value)

        # Bullet point without bold: "- Item"
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")

        # Markdown bold in regular line: "**text** more"
        elif "**" in line:
            p = doc.add_paragraph()
            while "**" in line:
                pre, rest = line.split("**", 1)
                bold_text, line = rest.split("**", 1)
                p.add_run(pre)
                p.add_run(bold_text).bold = True
            if line:
                p.add_run(line)

        # Regular paragraph
        else:
            doc.add_paragraph(line)

    return doc


async def fn_analyze_document(file:bytes,session_id) -> ApiResponse:
    try:
        # Step 1: Analyze document using Azure Document Intelligence
        analysis_result = await fn_extract_table_data(file)
        if analysis_result.statusCode != 200:
            return ApiResponse(statusCode=analysis_result.statusCode, message=analysis_result.message)

        document_data = analysis_result.data

        # Step 2: Extract identifiers using OpenAI (only document_title_phrase now)
        identifier_result = await fn_extract_identifiers_with_openai(document_data)
        if identifier_result.statusCode != 200:
                return ApiResponse(
                statusCode=identifier_result.statusCode,
                message=identifier_result.message,
                data=None
            )


        # Step 3: Add identifiers to document data and select primary ID (which is now document_title_phrase)
        identifiers_parsed = identifier_result.data.get("identifiers", {})
        document_data["identifiers"] = identifiers_parsed


        # Step 4: Save tables and identifiers to Excel
        # Pass the entire identifiers_parsed dict
        excel_file_path = await fn_save_tables_to_excel(document_data, identifiers_parsed)
        if detect_language_from_user_history(session_id) == 'en':
            return ApiResponse(
                statusCode=200,
                message="Table Extracted successfully",
                data={
                    "file_path": excel_file_path,
                    "overall_confidence": document_data.get("overall_confidence"), 
                    "identifiers": identifiers_parsed # Will only contain document_title_phrase
                }
            )
        else:
            return ApiResponse(
                statusCode=200,
                message="A táblázat sikeresen ki lett nyerve",
                data={
                    "file_path": excel_file_path,
                    "overall_confidence": document_data.get("overall_confidence"),
                    "identifiers": identifiers_parsed  # Csak a dokumentum címe szerepel benne
                }
            )

    except Exception as e:
        return ApiResponse(statusCode=500, message=str(e), data=None)


async def fn_extract_table_data(file: bytes) -> ApiResponse:
    try:

        # Upload to blob and get SAS URL
        sas_url = fn_upload_blob_and_get_sas(file)

        poller = document_client.begin_analyze_document_from_url(
            model_id="prebuilt-document",
            document_url=sas_url
        )
        result = poller.result()

        data = {
            "pages": [],
            "tables": [],
            "key_value_pairs": [],
            "paragraphs": [],
            "overall_confidence": None
        }

        confidences = []

        # Process paragraphs
        if hasattr(result, 'paragraphs'):
            for para in result.paragraphs:
                para_data = {
                    "content": para.content,
                    "role": getattr(para, 'role', None)
                }

                if hasattr(para, 'bounding_regions') and para.bounding_regions:
                    para_data["bounding_region"] = {
                        "page_number": para.bounding_regions[0].page_number,
                        "polygon": [point for point in para.bounding_regions[0].polygon]
                    }

                if hasattr(para, 'confidence') and para.confidence is not None:
                    para_data["confidence"] = para.confidence
                    confidences.append(para.confidence)

                data["paragraphs"].append(para_data)

        # Process key-value pairs
        if hasattr(result, 'key_value_pairs'):
            for kvp in result.key_value_pairs:
                confidence = kvp.confidence
                kvp_data = {
                    "key": kvp.key.content if kvp.key else None,
                    "value": kvp.value.content if kvp.value else None
                }

                if confidence is not None:
                    kvp_data["confidence"] = confidence
                    confidences.append(confidence)

                if kvp.key and hasattr(kvp.key, 'bounding_regions') and kvp.key.bounding_regions:
                    kvp_data["key_bounding_region"] = {
                        "page_number": kvp.key.bounding_regions[0].page_number,
                        "polygon": [point for point in kvp.key.bounding_regions[0].polygon]
                    }

                if kvp.value and hasattr(kvp.value, 'bounding_regions') and kvp.value.bounding_regions:
                    kvp_data["value_bounding_region"] = {
                        "page_number": kvp.value.bounding_regions[0].page_number,
                        "polygon": [point for point in kvp.value.bounding_regions[0].polygon]
                    }

                data["key_value_pairs"].append(kvp_data)

        # Process tables
        for table_idx, table in enumerate(result.tables):
            table_data = {
                "table_id": table_idx,
                "row_count": table.row_count,
                "column_count": table.column_count,
                "cells": [],
                "bounding_region": None
            }

            if hasattr(table, 'bounding_regions') and table.bounding_regions:
                table_data["bounding_region"] = {
                    "page_number": table.bounding_regions[0].page_number,
                    "polygon": [point for point in table.bounding_regions[0].polygon]
                }

            for cell in table.cells:
                cell_entry = {
                    "content": cell.content.strip() if cell.content else "",
                    "row_index": cell.row_index,
                    "column_index": cell.column_index,
                    "row_span": getattr(cell, 'row_span', 1),
                    "column_span": getattr(cell, 'column_span', 1),
                    "is_header": getattr(cell, 'is_header', False),
                    "is_footer": getattr(cell, 'is_footer', False)
                }

                if hasattr(cell, 'bounding_regions') and cell.bounding_regions:
                    cell_entry["bounding_region"] = {
                        "page_number": cell.bounding_regions[0].page_number,
                        "polygon": [point for point in cell.bounding_regions[0].polygon]
                    }

                if hasattr(cell, "confidence") and cell.confidence is not None:
                    cell_entry["confidence"] = cell.confidence
                    confidences.append(cell.confidence)

                table_data["cells"].append(cell_entry)

            data["tables"].append(table_data)

        # Process pages (just dimensions for now)
        for page in result.pages:
            data["pages"].append({
                "page_number": page.page_number,
                "width": page.width,
                "height": page.height,
                "unit": page.unit
            })
        
        for page in result.pages:
            for word in page.words:
                if hasattr(word, 'confidence') and word.confidence is not None:
                    confidences.append(word.confidence)
        # Average confidence from all valid elements
        valid_confidences = [c for c in confidences if isinstance(c, (float, int)) and 0.0 <= c <= 1.0]
        if valid_confidences:
            data["overall_confidence"] = round(sum(valid_confidences) / len(valid_confidences), 4)
        else:
            data["overall_confidence"] = None

        # Add metadata
        data["extraction_metadata"] = {
            "model_used": "prebuilt-document",
            "total_tables": len(data["tables"]),
            "total_paragraphs": len(data["paragraphs"]),
            "total_key_value_pairs": len(data["key_value_pairs"]),
            "extraction_timestamp": datetime.now().isoformat()
        }

        return ApiResponse(
            statusCode=200,
            message="Document analyzed successfully.",
            data=data
        )

    except Exception as e:
        return ApiResponse(statusCode=500, message=f"Error extracting table data: {str(e)}", data={})
    



async def fn_extract_identifiers_with_openai(data: dict) -> ApiResponse:
    try:
        # Prepare input context for OpenAI
        extraction_input = {
            "pages": data.get("pages", []),
            "paragraphs": data.get("paragraphs", []),
            "key_value_pairs": data.get("key_value_pairs", [])
        }

        prompt = """
            You are a document analysis assistant. Your task is to extract the most prominent document identifier from a scanned or digital invoice. This identifier is typically a combination of:

            - Company name (e.g., RENAULT TRUCKS)
            - Document type (e.g., Facture, Invoice, Számla)
            - A numeric ID (e.g., 8016184969)

            This identifier may appear in English, or Hungarian (e.g., "Számlaszám", "Ügyfélszám", etc.).

            **Only** return the full identifier phrase exactly as it appears in the document.

            If nothing is clearly identifiable, return `null`.

            Strictly respond in this JSON format:
            {
            "document_title_phrase": "RENAULT TRUCKS Facture: 8016184969"
            }
            or:
            {
            "document_title_phrase": null
            }
        """

        response = azure_client.chat.completions.create(
            model=AZURE_DEPLOYMENT_NAME,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": f"Document data:\n{json.dumps(extraction_input)[:4000]}"}
            ],
            temperature=0.2
        )

        content = response.choices[0].message.content.strip()

        if content.startswith("```json"):
            content = content[7:]
        if content.endswith("```"):
            content = content[:-3]

        try:
            identifiers = json.loads(content)
            return ApiResponse(
                statusCode=200,
                message="Unique identifier extracted successfully",
                data={"identifiers": identifiers}
            )
        except json.JSONDecodeError:
            pass  # fallback below

        # Fallback: Try to extract with regex from paragraph or page lines
        import re
        all_text = []

        for para in extraction_input.get("paragraphs", []):
            all_text.append(para.get("content", ""))

        for page in extraction_input.get("pages", []):
            for line in page.get("lines", []):
                all_text.append(line.get("content", ""))

        joined = " ".join(all_text)
        match = re.search(r"(RENAULT\s+TRUCKS.*?)\s+(Facture[:\s]+[0-9]+)", joined, re.IGNORECASE)
        if match:
            phrase = f"{match.group(1).strip()} {match.group(2).strip()}"
            return ApiResponse(
                statusCode=200,
                message="Unique identifier extracted via fallback",
                data={"identifiers": {"document_title_phrase": phrase}}
            )

        return ApiResponse(
            statusCode=200,
            message="No unique identifier found",
            data={"identifiers": {"document_title_phrase": None}}
        )

    except Exception as e:
        return ApiResponse(
            statusCode=500,
            message=f"Error extracting identifier: {str(e)}",
            data={"identifiers": {"document_title_phrase": None}}
        )


async def fn_save_tables_to_excel(data: dict, identifiers: dict) -> str:
    """This function is used to save extract data into excel and excel file is uploaded to blob"""
    temp_file_path = None
    try:
        all_tables_raw = data.get("tables", [])
        unique_identifier = identifiers.get("document_title_phrase", "N/A")

        processed_tables = []

        for table_idx, table_raw in enumerate(all_tables_raw):
            if table_raw.get("row_count", 0) == 0:
                continue

            max_row = table_raw["row_count"]
            max_col = table_raw["column_count"]

            table_grid = [["" for _ in range(max_col)] for _ in range(max_row)]

            for cell in table_raw["cells"]:
                row_idx = cell["row_index"]
                col_idx = cell["column_index"]
                content = cell.get("content", "")
                cleaned = content.replace(":selected:", "").replace(":unselected:", "").strip()
                cleaned = ' '.join(cleaned.split())
                table_grid[row_idx][col_idx] = cleaned

            if not table_grid or not any(any(row) for row in table_grid):
                continue

            headers = table_grid[0]
            rows = table_grid[1:] if len(table_grid) > 1 else []

            df = pd.DataFrame(rows, columns=headers)
            df = df.dropna(how='all')
            df = df.dropna(axis=1, how='all')

            if not df.empty:
                processed_tables.append({
                    "table_id": table_idx,
                    "data_df": df
                })

        # Create temp Excel file
        with NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_file:
            temp_file_path = temp_file.name

        with pd.ExcelWriter(temp_file_path, engine="openpyxl") as writer:
            for table in processed_tables:
                df = table["data_df"]
                table_id = table["table_id"]
                sheet_name = f"Table_{table_id + 1}"

                metadata_rows = [
                    ["Unique Identifier", unique_identifier],
                    ["Extraction Date", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
                    ["", ""]
                ]

                metadata_df = pd.DataFrame(metadata_rows, columns=["Field", "Value"])
                metadata_df.to_excel(writer, sheet_name=sheet_name, index=False, header=False, startrow=0)

                start_row = len(metadata_rows) + 1
                df.to_excel(writer, sheet_name=sheet_name, index=False, startrow=start_row)

                # Apply formatting
                worksheet = writer.sheets[sheet_name]

                # Style metadata
                header_font = Font(bold=True, color="FFFFFF")
                header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                for row in range(1, len(metadata_rows) + 1):
                    cell = worksheet.cell(row=row, column=1)
                    if cell.value:
                        cell.font = header_font
                        cell.fill = header_fill

                # Style table headers
                table_header_font = Font(bold=True)
                table_header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
                header_row = start_row + 1
                for col in range(1, len(df.columns) + 1):
                    cell = worksheet.cell(row=header_row, column=col)
                    cell.font = table_header_font
                    cell.fill = table_header_fill
                    cell.alignment = Alignment(horizontal="center")


        # Upload and return URL
        sas_url = fn_upload_blob_and_get_sas(temp_file_path)
        return sas_url

    except Exception as e:
        raise Exception(f"Error saving to Excel: {str(e)}")
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.unlink(temp_file_path)


def fn_upload_blob_and_get_sas(file_input, container_name=CONTAINER_NAME, filename: str = None) -> str:
    """ This function used to upload files into blob storage"""
    temp_file_path = None
    try:
        # Case 1: file_input is a file path
        if isinstance(file_input, str):
            file_path = file_input
            blob_name = f"tables/{os.path.basename(file_path)}"

        # Case 2: file_input is raw bytes
        elif isinstance(file_input, bytes):
            # Use provided filename or default
            if not filename:
                filename = f"uploaded_{uuid.uuid4().hex}.xlsx"
            temp_file = NamedTemporaryFile(delete=False, suffix=".xlsx")
            temp_file.write(file_input)
            temp_file.close()
            file_path = temp_file.name
            blob_name = f"tables/{filename}"
            temp_file_path = file_path  # For cleanup
        else:
            raise ValueError("Invalid file_input: must be file path or bytes")

        # Upload to blob
        blob_service = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        blob_client = blob_service.get_blob_client(container=container_name, blob=blob_name)

        with open(file_path, "rb") as f:
            blob_client.upload_blob(f, overwrite=True)

        # Generate SAS
        sas_token = generate_blob_sas(
            account_name=blob_service.account_name,
            container_name=container_name,
            blob_name=blob_name,
            account_key=AZURE_STORAGE_ACCOUNT_KEY,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.now(timezone.utc) + timedelta(hours=1)
        )

        return f"https://{blob_service.account_name}.blob.core.windows.net/{container_name}/{blob_name}?{sas_token}"

    except Exception as e:
        raise Exception(f"Blob upload failed: {str(e)}")
    finally:
        # Clean up temp file if created
        if temp_file_path and os.path.exists(temp_file_path):
            os.unlink(temp_file_path)

def detect_language_from_user_history(session_id: str) -> str:
    """This method is used to Determine the language used in the latest 
        message corresponding to a particular session ID. """
    try:
        history = CONVERSATION_HISTORY.get(session_id, {}).get("History", [])
        for msg in reversed(history):
            if msg.get("role") == "user" and msg.get("content"):
                return detect(msg["content"])
    except Exception as e:
        print(f"Language detection failed: {e}")
    return "en"